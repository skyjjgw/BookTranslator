from docx import Document
import sys

def read_docx(file_path):
    try:
        doc = Document(file_path)
        print("=== Paragraphs ===")
        for i, para in enumerate(doc.paragraphs):
            if para.text.strip():
                print(f"[{i}] {para.text}")
        
        print("\n=== Tables ===")
        for i, table in enumerate(doc.tables):
            print(f"Table {i}:")
            for r, row in enumerate(table.rows):
                row_data = []
                for cell in row.cells:
                    row_data.append(cell.text.strip())
                print(f"  Row {r}: {' | '.join(row_data)}")
                
    except Exception as e:
        print(f"Error: {e}")

if __name__ == "__main__":
    read_docx(r"C:\Users\skyjj\Desktop\就业+考研\就业\简历.docx")