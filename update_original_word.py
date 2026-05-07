from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH

def update_resume_docx():
    doc = Document()
    
    # 设置全局字体（中文：微软雅黑，英文：Arial）
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(10.5)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    
    # 调整页边距（让简历显得更紧凑，一页纸排版）
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

    # ================= 1. 头部：个人信息 =================
    p_name = doc.add_paragraph()
    p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_name = p_name.add_run("你的姓名")
    run_name.bold = True
    run_name.font.size = Pt(18)
    
    p_contact = doc.add_paragraph()
    p_contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_contact.add_run("手机：188-XXXX-XXXX  |  邮箱：your_email@xxx.com  |  微信：YourWeChat\n")
    p_contact.add_run("GitHub：github.com/YourName  |  求职意向：AI智能体开发 / 算法工程师")

    # ================= 2. 教育背景 =================
    doc.add_heading('教育背景', level=1)
    
    p_edu = doc.add_paragraph()
    run_edu1 = p_edu.add_run("某某大学    计算机科学与技术    本科/硕士")
    run_edu1.bold = True
    run_edu2 = p_edu.add_run(" " * 60 + "201X.09 - 202X.06")
    
    p_edu_details = doc.add_paragraph(style='List Bullet')
    p_edu_details.add_run("核心课程：数据结构、算法分析、人工智能、操作系统、计算机网络等")
    p_edu_details2 = doc.add_paragraph(style='List Bullet')
    p_edu_details2.add_run("荣誉奖项：连续两年获得国家励志奖学金、ACM省级二等奖（可选）")

    # ================= 3. 专业技能 =================
    doc.add_heading('专业技能', level=1)
    
    skills = [
        ("编程语言：", "熟练掌握 Python，熟悉 C++ / Java / Go 等编程语言。"),
        ("Agent 与大模型开发：", "熟练使用 LangChain、LangGraph 框架开发 AI Agent；熟悉 COZE、AutoGen、MetaGPT 等智能体搭建与多智能体协同框架；掌握 Prompt Engineering 技巧。"),
        ("RAG 与微调：", "掌握 RAG（检索增强生成）架构与 RAGFlow 本地化知识库搭建；了解 LLaMA3 等开源大模型的本地化部署、量化与微调（Fine-tuning）技术。"),
        ("多模态与算法：", "熟悉 CV/NLP 经典大模型与算法（如 GPT 系列、Diffusion、SAM、YOLO-World 等）；熟练使用 Pandas、pdfplumber 等进行多模态数据（文档、表格、图像）的提取与清洗。")
    ]
    
    for title, desc in skills:
        p_skill = doc.add_paragraph(style='List Bullet')
        p_skill.add_run(title).bold = True
        p_skill.add_run(desc)

    # ================= 4. 项目经历 =================
    doc.add_heading('项目经历', level=1)
    
    p_proj = doc.add_paragraph()
    run_proj1 = p_proj.add_run("智能文档翻译 Agent 系统 | 核心开发者")
    run_proj1.bold = True
    run_proj2 = p_proj.add_run(" " * 50 + "2023.10 - 至今")
    
    p_tech = doc.add_paragraph()
    p_tech.add_run("技术栈：Python, LangChain, DeepSeek/OpenAI API, Prompt Engineering, pdfplumber, ReportLab").italic = True

    bullets = [
        ("独立设计", "基于 LangChain 的文档翻译 Agent 架构，集成 DeepSeek/OpenAI 模型，构建涵盖解析、翻译与排版重建的全自动化流水线，使单本文档翻译耗时下降 95%。"),
        ("主导", "PDF 结构化解析，精准分离文本与表格；运用 Prompt Engineering 约束大模型严格输出 JSON，实现复杂表格 100% 无损翻译与跨语种格式还原。"),
        ("研发", "动态渲染引擎，解决跨语种字体乱码痛点，支持高保真 PDF 与 Markdown 多端输出；采用单例模式重构配置中心，大幅提升系统可扩展性与健壮性。")
    ]
    
    for bold_text, normal_text in bullets:
        p_bullet = doc.add_paragraph(style='List Bullet')
        p_bullet.add_run(bold_text).bold = True
        p_bullet.add_run(normal_text)

    # ================= 5. 个人亮点 =================
    doc.add_heading('个人亮点', level=1)
    
    highlights = [
        ("技术热情：", "热爱开源，对前沿大模型技术（如 MOE 多专家系统、多模态融合）保持高度关注并具备极强的自驱学习能力。"),
        ("英语能力：", "CET-6，能流畅阅读英文官方文档与顶会 Paper。")
    ]
    
    for title, desc in highlights:
        p_hl = doc.add_paragraph(style='List Bullet')
        p_hl.add_run(title).bold = True
        p_hl.add_run(desc)

    # 修改所有标题的样式
    for paragraph in doc.paragraphs:
        if paragraph.style.name.startswith('Heading 1'):
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(0, 0, 0)
                run.font.size = Pt(14)
                run.bold = True
            # 添加下边框的效果（用下划线替代）
            p_border = doc.add_paragraph()
            p_border.add_run("_" * 75)
            # 把下划线移到标题下方
            paragraph._p.addnext(p_border._p)

    output_path = r'D:\BaiduNetdiskDownload\BookTranslator\AI_Agent_Resume_Project.docx'
    doc.save(output_path)
    print(f"原始简历 Word 文档已覆盖更新：{output_path}")

if __name__ == '__main__':
    update_resume_docx()