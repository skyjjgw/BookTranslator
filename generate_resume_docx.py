import os
import subprocess
import sys

def install_and_import(package):
    try:
        __import__(package)
    except ImportError:
        print(f"Installing {package}...")
        subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])

install_and_import('docx')
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

def generate_word():
    doc = Document()
    
    # 设置中文字体
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Microsoft YaHei'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    doc.add_heading('简历项目描述 - AI智能体（AI Agent）岗位', 0)

    p_header = doc.add_paragraph()
    p_header.add_run('项目名称：').bold = True
    p_header.add_run('智能文档翻译 Agent 系统\n')
    p_header.add_run('担任角色：').bold = True
    p_header.add_run('核心开发者\n')
    p_header.add_run('技术栈：').bold = True
    p_header.add_run('Python, LangChain, DeepSeek/OpenAI API, Prompt Engineering, pdfplumber, ReportLab')

    doc.add_heading('核心贡献与成果 (STAR法则)', level=2)

    bullets = [
        "独立设计基于 LangChain 的文档翻译 Agent 架构，集成 DeepSeek/OpenAI 模型，构建涵盖解析、翻译与排版重建的全自动化流水线，使单本文档翻译耗时下降 95%。",
        "主导 PDF 结构化解析，精准分离文本与表格；运用 Prompt Engineering 约束大模型严格输出 JSON，实现复杂表格 100% 无损翻译与跨语种格式还原。",
        "研发动态渲染引擎，解决跨语种字体乱码痛点，支持高保真 PDF 与 Markdown 多端输出；采用单例模式重构配置中心，大幅提升系统可扩展性与健壮性。"
    ]

    for bullet in bullets:
        doc.add_paragraph(bullet, style='List Bullet')
        
    doc.add_heading('ATS 关键词匹配解析 (面试备用)', level=2)
    ats_text = "AI/大模型方向：AI Agent（智能体）、LangChain、DeepSeek/OpenAI API、Prompt Engineering（提示词工程）、JSON 结构化输出\n工程/数据方向：自动化流水线、结构化解析、单例模式、多模态/多端输出"
    doc.add_paragraph(ats_text)

    output_path = r'D:\BaiduNetdiskDownload\BookTranslator\AI_Agent_Resume_Project.docx'
    doc.save(output_path)
    print(f"Word文档已成功生成：{output_path}")

if __name__ == '__main__':
    generate_word()
