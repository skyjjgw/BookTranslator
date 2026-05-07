from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH

def update_existing_resume():
    # 打开原始简历
    doc = Document(r"C:\Users\skyjj\Desktop\就业+考研\就业\简历.docx")
    
    # 定义中文字体设置函数
    def set_font(run, font_name='Microsoft YaHei', size=10.5):
        run.font.name = font_name
        run.font.size = Pt(size)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

    # =========================================================
    # 第一步：修改求职意向（段落 4-5 左右）
    # =========================================================
    for para in doc.paragraphs:
        if "意向岗位" in para.text:
            para.clear()
            run = para.add_run("意向岗位: AI智能体开发 / 算法工程师" + " " * 30 + "求职类型: 实习")
            set_font(run, size=11)
            break

    # =========================================================
    # 第二步：在“在校经历”之前，插入全新的“专业技能”模块
    # =========================================================
    # 找到“在校经历”标题的段落索引
    insert_index = -1
    for i, para in enumerate(doc.paragraphs):
        if "在校经历" in para.text or "项目经历" in para.text:
            insert_index = i
            break
            
    if insert_index != -1:
        # 在“在校经历”上方插入
        target_para = doc.paragraphs[insert_index]
        
        # 插入标题
        p_skill_title = target_para.insert_paragraph_before()
        run_title = p_skill_title.add_run("   专业技能")
        set_font(run_title, size=14)
        run_title.bold = True
        
        # 插入技能点
        skills = [
            ("编程语言：", "熟练掌握 Python，熟悉 C++ / Java / Go 等编程语言。"),
            ("Agent 与大模型开发：", "熟练使用 LangChain、LangGraph 框架开发 AI Agent；熟悉 COZE、AutoGen、MetaGPT 等智能体搭建与多智能体协同框架；掌握 Prompt Engineering 技巧。"),
            ("RAG 与微调：", "掌握 RAG（检索增强生成）架构与 RAGFlow 本地化知识库搭建；了解 LLaMA3 等开源大模型的本地化部署、量化与微调（Fine-tuning）技术。"),
            ("多模态与算法：", "熟悉 CV/NLP 经典大模型与算法（如 GPT 系列、Diffusion、SAM、YOLO-World 等）；熟练使用 Pandas、pdfplumber 等进行多模态数据（文档、表格、图像）的提取与清洗。")
        ]
        
        for title, desc in skills:
            p_skill = target_para.insert_paragraph_before()
            # 模拟列表样式，添加缩进
            p_skill.paragraph_format.left_indent = Inches(0.2)
            run_b = p_skill.add_run("• " + title)
            run_b.bold = True
            set_font(run_b)
            run_n = p_skill.add_run(desc)
            set_font(run_n)

    # =========================================================
    # 第三步：在“在校经历”下方插入最新的 AI Agent 项目
    # =========================================================
    # 我们将其插在“视桥智导”或者最前面
    exp_index = -1
    for i, para in enumerate(doc.paragraphs):
        if "2024.05-2025.6" in para.text and "创新训练计划项目" in para.text:
            exp_index = i
            break
            
    if exp_index != -1:
        target_para = doc.paragraphs[exp_index]
        
        p_proj_title = target_para.insert_paragraph_before()
        run_t1 = p_proj_title.add_run("2023.10 - 至今" + " " * 30 + "智能文档翻译 Agent 系统\n")
        run_t1.bold = True
        set_font(run_t1, size=11)
        run_t2 = p_proj_title.add_run("核心开发者\n")
        set_font(run_t2, size=11)
        run_t3 = p_proj_title.add_run("技术栈：Python, LangChain, DeepSeek/OpenAI API, Prompt Engineering, pdfplumber, ReportLab")
        run_t3.italic = True
        set_font(run_t3, size=10)

        bullets = [
            ("独立设计：", "基于 LangChain 的文档翻译 Agent 架构，集成 DeepSeek/OpenAI 模型，构建涵盖解析、翻译与排版重建的全自动化流水线，使单本文档翻译耗时下降 95%。"),
            ("主导开发：", "PDF 结构化解析，精准分离文本与表格；运用 Prompt Engineering 约束大模型严格输出 JSON，实现复杂表格 100% 无损翻译与跨语种格式还原。"),
            ("架构优化：", "研发动态渲染引擎，解决跨语种字体乱码痛点，支持高保真 PDF 与 Markdown 多端输出；采用单例模式重构配置中心，大幅提升系统可扩展性与健壮性。")
        ]
        
        for bold_text, normal_text in bullets:
            p_bullet = target_para.insert_paragraph_before()
            p_bullet.paragraph_format.left_indent = Inches(0.2)
            run_b = p_bullet.add_run("• " + bold_text)
            run_b.bold = True
            set_font(run_b)
            run_n = p_bullet.add_run(normal_text)
            set_font(run_n)
            
    # =========================================================
    # 第四步：保存文件
    # =========================================================
    output_path = r"C:\Users\skyjj\Desktop\就业+考研\就业\简历.docx"
    doc.save(output_path)
    print(f"简历已成功更新并覆盖：{output_path}")

if __name__ == '__main__':
    update_existing_resume()